"""
Generate the project report (.docx) for IN 22-21 submission.

Re-run after edits:  python build_report.py
Output:              reports/TND_Intrinsic_Value_Report.docx

Requires:  pip install python-docx
"""
from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "reports" / "TND_Intrinsic_Value_Report.docx"


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for lvl, sz, color in [(1, 18, "1F4E79"), (2, 14, "1F4E79"), (3, 12, "2E75B6")]:
        s = styles[f"Heading {lvl}"]
        s.font.name = "Calibri"
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(28)
    title.font.color.rgb = RGBColor.from_string("1F4E79")

    # Code style
    if "Mono" not in [st.name for st in styles]:
        mono = styles.add_style("Mono", WD_STYLE_TYPE.PARAGRAPH)
        mono.font.name = "Consolas"
        mono.font.size = Pt(9)
        mono.paragraph_format.left_indent = Cm(0.5)
        mono.paragraph_format.space_before = Pt(4)
        mono.paragraph_format.space_after = Pt(4)


def add_para(doc, text: str, *, bold=False, italic=False, style="Normal"):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it)


def add_code_block(doc, text: str):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph(style="Mono")
        p.add_run(line)


def add_kv_table(doc, rows, header=("Field", "Value")):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    h = t.rows[0].cells
    h[0].text = header[0]
    h[1].text = header[1]
    for c in h:
        for r in c.paragraphs[0].runs:
            r.bold = True
        _set_cell_shading(c, "1F4E79")
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = RGBColor.from_string("FFFFFF")
    for k, v in rows:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v
    return t


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    _setup_styles(doc)

    # ------------------------------ Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\n")

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Real-Time Intrinsic TND\nValuation Model")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("A basket-based baseline with Kalman-filtered liquidity adjustment\nfor the USD/TND exchange rate")
    rs.italic = True
    rs.font.size = Pt(13)
    rs.font.color.rgb = RGBColor.from_string("595959")

    doc.add_paragraph("\n\n")

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # NOTE: edit AUTHORS here, or in Word after generation. The placeholder is
    # intentionally explicit so it cannot be missed at proof-reading time.
    AUTHORS = "« Replace with author names »"
    meta_rows = [
        ("Course", "IN 22-21 — Time Series Analysis"),
        ("Instructor", "Dr Eymen Errais"),
        ("Authors", AUTHORS),
        ("Date", date.today().strftime("%B %Y")),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v
        for c in meta.rows[i].cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ------------------------------ Table of contents (placeholder)
    doc.add_heading("Table of Contents", level=1)
    add_para(doc, "Right-click → Update Field in Word to populate. (Placeholder ToC.)", italic=True)
    doc.add_paragraph()
    toc_items = [
        "1. Executive Summary",
        "2. Background and Problem Statement",
        "3. Data",
        "4. Model Framework",
        "5. Basket Weight Estimation",
        "6. Stochastic Liquidity Adjustment",
        "7. Interbank Rate Nowcasting under Data Lag",
        "8. Real-Time / Intraday Operation",
        "9. Performance Evaluation and Backtesting",
        "10. Limitations and Future Work",
        "11. Conclusion",
        "Appendix A — Repository Structure",
        "Appendix B — Database Schema",
        "Appendix C — Key Equations",
    ]
    for t in toc_items:
        doc.add_paragraph(t)
    doc.add_page_break()

    # ------------------------------ 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    add_para(
        doc,
        "This report presents a real-time intrinsic-value model for the USD/TND exchange rate. "
        "The model combines a fundamental basket-based baseline — driven by global FX movements "
        "in EUR/USD, GBP/USD, and USD/JPY — with a stochastic adjustment that captures local "
        "Tunisian FX-market frictions, primarily liquidity pressure between the official BCT "
        "fixing and the lagged interbank (IB) rate.",
    )
    add_para(
        doc,
        "The implementation is delivered as a zero-cost Python pipeline with a SQLite store, an "
        "Excel reporting layer, and a Streamlit dashboard. Two operational modes are supported: a "
        "daily batch (orchestrated by run_pipeline.py) and a true intraday loop (run_realtime.py) "
        "that ticks at minute resolution and emits a continuously updated intrinsic value.",
    )
    add_para(doc, "Key contributions:", bold=True)
    add_bullets(doc, [
        "Rolling 90-day OLS estimation of basket weights with Newey-West (HAC) standard errors, t-stats and p-values.",
        "AR(1) + Kalman filter on the IB-Fix spread with joint maximum-likelihood calibration of (c, φ, Q, R).",
        "Real-time intraday engine using free 1-minute FX bars and an anchored basket return; ±2σ confidence band from the Kalman state covariance.",
        "Walk-forward backtest with MAE / RMSE / MAPE / directional accuracy, Diebold-Mariano test vs random-walk, Ljung-Box on residuals, ADF / KPSS on the spread.",
        "Two-fixing (AM / PM) schema with documented choice of which fixing anchors the basket vs the spread.",
        "Read-only FastAPI dashboard with premium / discount narrative, live intraday section, and an embedded LLM analyst (Claude / Groq) that answers user questions grounded on the current market snapshot.",
        "Macro-overlay regression with FRED / Yahoo Finance covariates (Brent, broad USD, VIX) — finds the broad-dollar index significant at p<0.001 as a driver of the IB-Fix spread.",
        "Operational channels: Excel report (5 sheets), Word document, Streamlit (legacy) and FastAPI (current) dashboards, optional Telegram push notification at end of the daily run.",
    ])
    add_para(doc, "Headline empirical finding:", bold=True)
    add_para(
        doc,
        "On out-of-sample data, the model achieves high R² in levels (≈0.97) but is "
        "outperformed by a one-day random-walk baseline in absolute error (DM p≈0). "
        "Combined with strongly autocorrelated residuals (Ljung-Box p≈0) and a "
        "borderline-non-stationary spread, this indicates that local Tunisian FX "
        "dynamics carry structure not captured by a linear basket + AR(1) Kalman. "
        "The honest interpretation is that BCT actively manages the fixing rather "
        "than letting it freely track the basket — none of the basket coefficients "
        "is statistically significant after Newey-West correction on the full "
        "sample. This motivates the regime-switching and macro-overlay extensions "
        "discussed in §10.",
    )

    # ------------------------------ 2. Background
    doc.add_heading("2. Background and Problem Statement", level=1)
    add_para(
        doc,
        "The Tunisian Dinar (TND) exchange rate against the US Dollar (USD) is officially "
        "determined by the Central Bank of Tunisia (BCT) and is published as an official fixing "
        "rate twice daily — morning and evening. These fixings serve as benchmarks and are "
        "broadly aligned with global currency movements, particularly EUR/USD, GBP/USD, and "
        "USD/JPY. Conceptually, the intrinsic value of USD/TND can be expressed as a weighted "
        "basket of these major currency pair returns:",
    )
    add_code_block(doc, "USD/TND = (Latest fixing) × (w1·%ΔEUR/USD + w2·%ΔGBP/USD + w3·%ΔUSD/JPY)")
    add_para(
        doc,
        "The actual market USD/TND rate, however, deviates from this basket-implied value due to "
        "local FX market conditions — liquidity scarcity, supply/demand imbalances, and "
        "regulatory frictions. The interbank (IB) rate, published with a one-day lag, is the "
        "best available proxy for the true local market price. The challenge is to model these "
        "deviations in real time, despite the IB data lag, to produce a continuously updating "
        "fair-value estimate.",
    )

    # ------------------------------ 3. Data
    doc.add_heading("3. Data", level=1)
    doc.add_heading("3.1 Sources", level=2)
    add_bullets(doc, [
        "Global FX (EUR/USD, GBP/USD, USD/JPY) — daily: ExchangeRate-API v6 (primary), exchangerate.host (fallback), frankfurter.app (fallback).",
        "Global FX intraday (1-minute bars): yfinance (yahoo finance, free, no API key).",
        "BCT official fixing (USD/TND): scraped from https://www.bct.gov.tn or supplied via the BCT_FIX_MID environment variable.",
        "Historical IB rates: data/ib.csv (seeded into fx_rates.ib_rate).",
        "Historical baseline FX + Fix: data/FX-CLEAN-Data.csv.",
    ])

    doc.add_heading("3.2 Conventions", level=2)
    add_kv_table(doc, [
        ("EUR/USD", "USD per 1 EUR"),
        ("GBP/USD", "USD per 1 GBP"),
        ("USD/JPY", "JPY per 1 USD"),
        ("fix_mid", "TND per 1 USD (BCT mid fixing)"),
        ("ib_rate", "TND per 1 USD (interbank, T-1 lag)"),
        ("Returns", "Log-returns: r_t = ln(P_t / P_{t-1})"),
        ("Spread", "spread_pub = ib_rate − fix_mid (units: TND)"),
    ])

    doc.add_heading("3.3 Persistence", level=2)
    add_para(doc, "All series are stored in a single SQLite database (data/tnd.db). The schema is reproduced in Appendix B.")

    # ------------------------------ 4. Framework
    doc.add_heading("4. Model Framework", level=1)
    add_para(
        doc,
        "The intrinsic value of USD/TND at time t is decomposed into two additive components:",
    )
    add_code_block(doc, "intrinsic_v2(t) = intrinsic_v1(t) + kf_state(t)\n"
                        "intrinsic_v1(t) = anchor_fix × exp( w_const + w_EUR·r_EUR + w_GBP·r_GBP + w_JPY·r_JPY )")
    add_bullets(doc, [
        "intrinsic_v1 — basket baseline. Anchored on the most recent BCT fixing and shocked by the "
        "log-return of the global basket since the anchor time.",
        "kf_state — stochastic local-liquidity adjustment. Filtered estimate of the IB-vs-Fix "
        "spread using an AR(1) state-space model.",
    ])
    add_para(
        doc,
        "Daily mode refits the basket weights and the Kalman parameters once per day from the "
        "lookback window (default 500 days). Intraday mode reuses the daily-fit parameters and "
        "advances the Kalman state forward at each tick — see Section 8.",
    )

    # ------------------------------ 5. Basket weights
    doc.add_heading("5. Basket Weight Estimation", level=1)
    doc.add_heading("5.1 Specification", level=2)
    add_para(doc, "The basket function f is estimated via ordinary least squares (OLS) on log-returns:")
    add_code_block(doc, "y_t  = ret_Fix = log(fix_mid_t / fix_mid_{t-1})\n"
                        "x_t  = [1, ret_EURUSD, ret_GBPUSD, ret_USDJPY]\n"
                        "y_t  = β · x_t + ε_t")
    add_para(doc, "Two estimators are produced:")
    add_bullets(doc, [
        "Full-sample OLS — used for goodness-of-fit reporting (R²).",
        "Rolling 90-day OLS — refit at each step; the latest row supplies weights to predict_today().",
    ])

    doc.add_heading("5.2 Methodology", level=2)
    add_bullets(doc, [
        "Closed-form solution via numpy.linalg.lstsq (Moore-Penrose pseudo-inverse).",
        "No regularization — the design matrix is well-conditioned at this dimensionality.",
        "Returns are demeaned implicitly via the intercept term w_const.",
        "Weights are not constrained to sum to 1; we report Σw as a stability diagnostic.",
    ])

    doc.add_heading("5.3 Results", level=2)
    # Auto-fill from a fresh fit_ols + rolling_weights on the current DB
    ols_dict = None
    last_w = None
    try:
        import sqlite3
        from clean_returns import load_and_clean
        from model import fit_ols as _fit, rolling_weights as _roll
        con = sqlite3.connect(str(ROOT / "data" / "tnd.db"))
        _df = load_and_clean(con, lookback_days=10_000)
        con.close()
        if not _df.empty and len(_df) >= 90:
            ols_dict = _fit(_df)
            _r = _roll(_df, 90)
            if not _r.empty and not pd.isna(_r.iloc[-1]["w_EURUSD"]):
                last_w = _r.iloc[-1].to_dict()
    except Exception as _e:
        ols_dict = None

    def _f(x, dp=6):
        if x is None or (isinstance(x, float) and (x != x)):
            return "—"
        try:
            return f"{float(x):.{dp}f}"
        except Exception:
            return str(x)

    if ols_dict is not None:
        add_para(doc, "Full-sample OLS with Newey-West HAC standard errors:", bold=True)
        add_kv_table(doc, [
            ("Coefficient", "Estimate · SE · t · p"),
            ("Intercept", f"{_f(ols_dict['intercept'])} · {_f(ols_dict['se_intercept'])} · {_f(ols_dict['t_intercept'], 3)} · {_f(ols_dict['p_intercept'], 4)}"),
            ("w_EURUSD", f"{_f(ols_dict['w_EURUSD'])} · {_f(ols_dict['se_w_EURUSD'])} · {_f(ols_dict['t_w_EURUSD'], 3)} · {_f(ols_dict['p_w_EURUSD'], 4)}"),
            ("w_GBPUSD", f"{_f(ols_dict['w_GBPUSD'])} · {_f(ols_dict['se_w_GBPUSD'])} · {_f(ols_dict['t_w_GBPUSD'], 3)} · {_f(ols_dict['p_w_GBPUSD'], 4)}"),
            ("w_USDJPY", f"{_f(ols_dict['w_USDJPY'])} · {_f(ols_dict['se_w_USDJPY'])} · {_f(ols_dict['t_w_USDJPY'], 3)} · {_f(ols_dict['p_w_USDJPY'], 4)}"),
            ("R² (full-sample)", _f(ols_dict["r_squared"], 5)),
            ("N observations", str(ols_dict.get("n_obs", "—"))),
            ("Newey-West lags", str(ols_dict.get("nw_lags", "—"))),
        ], header=("Field", "Value"))

        if last_w is not None:
            add_para(doc, "Latest 90-day rolling weights (used for live prediction):", bold=True)
            wsum = float(last_w["w_EURUSD"]) + float(last_w["w_GBPUSD"]) + float(last_w["w_USDJPY"])
            add_kv_table(doc, [
                ("w_EURUSD", _f(last_w["w_EURUSD"])),
                ("w_GBPUSD", _f(last_w["w_GBPUSD"])),
                ("w_USDJPY", _f(last_w["w_USDJPY"])),
                ("Σ weights", _f(wsum)),
                ("R² (rolling)", _f(last_w.get("R2"), 5)),
            ])
    else:
        add_para(doc, "[Run python run_pipeline.py to populate predictions, then re-run build_report.py]", italic=True)
    add_para(doc, "Diagnostics (to be added):", bold=True)
    add_bullets(doc, [
        "Newey-West heteroskedasticity- and autocorrelation-consistent standard errors.",
        "Weight-stability plot over the rolling window (chart in dashboard).",
        "Residual ACF and Ljung-Box test for serial correlation.",
        "ADF / KPSS stationarity tests on the return series.",
    ])

    # ------------------------------ 6. Liquidity
    doc.add_heading("6. Stochastic Liquidity Adjustment", level=1)
    doc.add_heading("6.1 Spread definition", level=2)
    add_para(doc, "The local liquidity premium is captured by the spread between the IB rate and the BCT fixing:")
    add_code_block(doc, "spread_t = ib_rate_t − fix_mid_t")

    doc.add_heading("6.2 AR(1) state-space model", level=2)
    add_para(doc, "The spread is modelled as a noisy AR(1) process with the standard Kalman filter recursion:")
    add_code_block(doc, "State:        x_t = c + φ · x_{t-1} + w_t,   w_t ~ N(0, Q)\n"
                        "Observation:  y_t = x_t + v_t,                v_t ~ N(0, R)")
    add_para(doc, "Parameters are estimated by joint maximum-likelihood:")
    add_bullets(doc, [
        "The Gaussian log-likelihood ℒ(c, φ, Q, R | y₁…y_T) is computed via the Kalman innovations form.",
        "Optimization: Nelder-Mead with multi-start; φ parameterized as tanh(z) to enforce |φ|<1 stationarity; Q and R as exp(·) for positivity.",
        "Steady-state covariance P_∞ obtained by fixed-point iteration of the Riccati equation.",
        "Pre-MLE OLS seed retained for fallback if SciPy is unavailable.",
    ])

    # Auto-fill MLE estimates from the current DB
    try:
        import sqlite3
        from clean_returns import load_and_clean
        from model import mle_kalman_ar1 as _mle
        con = sqlite3.connect(str(ROOT / "data" / "tnd.db"))
        _df = load_and_clean(con, lookback_days=10_000)
        con.close()
        if not _df.empty and _df["spread_pub"].dropna().shape[0] >= 30:
            m = _mle(_df["spread_pub"])
            add_para(doc, "MLE estimates on the in-sample spread series:", bold=True)
            add_kv_table(doc, [
                ("c (intercept)",        f"{m['c']:+.6f}"),
                ("φ (AR(1) coefficient)", f"{m['phi']:+.4f}"),
                ("Q (state variance)",   f"{m['Q']:.3e}"),
                ("R (obs. variance)",    f"{m['R']:.3e}"),
                ("Log-likelihood",       f"{m['loglik']:.2f}"),
                ("N observations",       str(m['n_obs'])),
                ("Method",               str(m['method'])),
            ])
    except Exception:
        pass

    doc.add_heading("6.3 Forward propagation under IB lag", level=2)
    add_para(
        doc,
        "Because the IB rate is published with a one-day lag, no observation update is available "
        "during the trading day. The intraday Kalman step therefore only propagates:",
    )
    add_code_block(doc, "x_pred = c + φ · x_{last_obs}\n"
                        "P_pred = φ² · P_{last_obs} + Q\n"
                        "σ_t   = √P_pred")
    add_para(
        doc,
        "Each new daily IB observation triggers a full Kalman update (filter step) and a refit of "
        "(c, φ, Q), invalidating the intraday parameter cache.",
    )

    # ------------------------------ 7. Nowcasting
    doc.add_heading("7. Interbank Rate Nowcasting under Data Lag", level=1)
    add_para(
        doc,
        "The BCT publishes the average interbank TND exchange rate with a one-day lag. The "
        "project specification asks for an algorithm that infers (\"nowcasts\") today's IB rate "
        "during the trading day. Our approach:",
    )
    add_bullets(doc, [
        "Anchor on yesterday's filtered Kalman state x_{T-1}.",
        "Propagate one step forward using the AR(1) transition: x_T_pred = c + φ · x_{T-1}.",
        "Surface the implied IB nowcast as: ib_nowcast(t) = intrinsic_v1(t) + x_T_pred.",
        "When new partial signals are available (e.g., bank indicative quotes), they can be "
        "fed as additional observations into the same Kalman update — the framework is open.",
    ])
    add_para(
        doc,
        "This produces a smooth, mean-reverting estimate of the IB rate that respects the AR(1) "
        "decay of the spread. The estimate uncertainty grows with √(P_pred), which the dashboard "
        "renders as a ±2σ band around the intrinsic value.",
    )

    # ------------------------------ 8. Real-Time
    doc.add_heading("8. Real-Time / Intraday Operation", level=1)
    doc.add_heading("8.1 Architecture", level=2)
    add_bullets(doc, [
        "realtime.tick(conn) — single tick: fetch quote → compute basket return → propagate Kalman → persist.",
        "run_realtime.py — loop runner with --interval and --once flags; default 60-second cadence.",
        "fx_intraday — raw quote table, minute resolution.",
        "intrinsic_intraday — emitted estimates table, including kf_state, kf_sigma, intrinsic_v1, intrinsic_v2.",
    ])

    doc.add_heading("8.2 Tick logic", level=2)
    add_code_block(doc, "1.  q   = yfinance(EURUSD=X, GBPUSD=X, JPY=X)            # latest 1-min bar\n"
                        "2.  A   = SELECT latest fx_rates row with fix_mid IS NOT NULL\n"
                        "3.  r_i = log(q_i / A_i)                                 # i ∈ {EUR, GBP, JPY}\n"
                        "4.  ret = w_const + w_EUR·r_EUR + w_GBP·r_GBP + w_JPY·r_JPY\n"
                        "5.  v1  = A.fix_mid · exp(ret)\n"
                        "6.  x   = c + φ · x_{last}                                # cached params\n"
                        "7.  σ   = √(φ² · P_{last} + Q)\n"
                        "8.  v2  = v1 + x\n"
                        "9.  INSERT INTO intrinsic_intraday(...)")

    doc.add_heading("8.3 On-board AI analyst (LLM assistant)", level=2)
    add_para(
        doc,
        "To make the model accessible to non-technical users, the dashboard ships "
        "with an embedded conversational assistant powered by the Anthropic Claude "
        "API. The assistant is rendered as a floating chat panel in the bottom-right "
        "of the SPA; clicking the \"Ask the analyst\" button opens a Perplexity-style "
        "thread with four pre-loaded prompts. The user can also type free-form "
        "questions.",
    )
    add_para(doc, "Architecture:", bold=True)
    add_bullets(doc, [
        "Endpoint: POST /api/chat on the FastAPI backend (serve.py).",
        "Request body: { message: string, history: [{role, content}, ...] }.",
        "Server-side: load the current snapshot from SQLite, embed it as JSON into the system prompt to ground answers in live data.",
        "Model: claude-haiku-4-5 — lowest-latency Claude model, ideal for an interactive widget.",
        "Conversation history capped at 16 turns to keep prompts small and costs negligible.",
        "Graceful fallback: if ANTHROPIC_API_KEY is unset, the endpoint returns an inline instruction message instead of failing.",
    ])
    add_para(doc, "System prompt design:", bold=True)
    add_para(
        doc,
        "The system prompt declares the assistant's role (on-board analyst), "
        "summarizes the model architecture (basket OLS with Newey-West HAC, "
        "AR(1) Kalman with MLE-calibrated parameters, MSK regime extension, "
        "backtest methodology) and the meaning of premium/discount. The current "
        "snapshot — latest fixings, IB rate, intrinsic value, premium in basis "
        "points, Kalman σ, basket weights, freshness flag — is appended to the "
        "system prompt at request time, so the assistant always answers about "
        "the current state of the market, not a frozen training snapshot.",
    )
    add_para(doc, "Default user prompts pre-loaded as suggestion chips:", bold=True)
    add_bullets(doc, [
        "\"What's the current premium telling us?\"",
        "\"How is the model built?\"",
        "\"Why does the model lose to a random walk?\"",
        "\"Explain the Markov-switching regimes.\"",
    ])
    add_para(doc, "Security:", bold=True)
    add_bullets(doc, [
        "The API key is read from the ANTHROPIC_API_KEY environment variable, never hard-coded.",
        "A .env loader is included in serve.py — a .env file at the repo root is auto-read at startup and is excluded from git via .gitignore.",
        "User messages are sent to Anthropic's API; no message data is persisted server-side.",
    ])

    doc.add_heading("8.4 Telegram notifications", level=2)
    add_para(
        doc,
        "The daily pipeline ends with an optional Telegram push (notify_telegram.py) "
        "that delivers the prediction summary to a configured chat. The notifier "
        "reads TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID from environment variables "
        "(or the .env file) and posts a single message containing the run date, "
        "the previous BCT fixing, intrinsic_v1, intrinsic_v2, the basket-implied "
        "return in percent, the Kalman spread component, and the model's R² for "
        "the day. If either credential is missing, the step is skipped silently — "
        "the pipeline never fails because of a missing notification channel.",
    )
    add_para(doc, "Operational contract:", bold=True)
    add_bullets(doc, [
        "Trigger: end of run_pipeline.py after a successful predict_today.",
        "Transport: Telegram Bot HTTP API (no extra dependency beyond requests).",
        "Format: short plain-text message, ~6 lines.",
        "Failure mode: best-effort; an HTTP error is logged but does not abort the pipeline.",
        "Security: tokens are never persisted to the repository or to SQLite.",
    ])

    doc.add_heading("8.5 Caching", level=2)
    add_para(
        doc,
        "Refitting the AR(1) + Kalman filter on every intraday tick is wasteful — the daily "
        "spread series only changes once per day. Parameters (c, φ, Q, x_{last}, P_{last}) are "
        "cached in a module-level dict keyed by the anchor date and only invalidated when a new "
        "daily observation arrives.",
    )

    # ------------------------------ 9. Backtesting
    doc.add_heading("9. Performance Evaluation and Backtesting", level=1)
    doc.add_heading("9.1 Methodology", level=2)
    add_bullets(doc, [
        "Walk-forward expanding-window split: train on [start, t], evaluate at t+1.",
        "Out-of-sample forecast: intrinsic_v2(t+1) computed from weights and Kalman state available at time t.",
        "Comparison benchmark: random-walk baseline (intrinsic_v2(t+1) = fix_mid(t)).",
    ])

    doc.add_heading("9.2 Metrics", level=2)
    # Auto-fill from reports/backtest_metrics.json if available
    bt_path = ROOT / "reports" / "backtest_metrics.json"
    if bt_path.exists():
        bt = json.loads(bt_path.read_text())
        s = bt.get("summary", {})
        st = bt.get("spread_stationarity", {})

        def f(x, dp=6):
            if x is None or (isinstance(x, float) and (x != x)):
                return "—"
            try:
                return f"{float(x):.{dp}f}"
            except Exception:
                return str(x)

        add_kv_table(doc, [
            ("N (out-of-sample)", str(s.get("N", "—"))),
            ("MAE", f(s.get("MAE"))),
            ("RMSE", f(s.get("RMSE"))),
            ("MAPE (%)", f(s.get("MAPE_pct"), 4)),
            ("Directional Accuracy (%)", f(s.get("Directional_Accuracy_pct"), 3)),
            ("Out-of-sample R²", f(s.get("OOS_R2"), 5)),
            ("MAE — Random Walk baseline", f(s.get("MAE_random_walk"))),
            ("RMSE — Random Walk baseline", f(s.get("RMSE_random_walk"))),
            ("Diebold-Mariano stat (RW − V2)", f(s.get("DM_stat"), 4)),
            ("Diebold-Mariano p-value", f(s.get("DM_pvalue_vs_RW"), 6)),
            ("Ljung-Box Q (lag 10)", f(s.get("LjungBox_Q_lag10"), 3)),
            ("Ljung-Box p (lag 10)", f(s.get("LjungBox_p_lag10"), 6)),
        ], header=("Metric", "Value"))

        add_para(doc, "Spread stationarity tests:", bold=True)
        add_kv_table(doc, [
            ("ADF statistic", f(st.get("ADF_stat"), 4)),
            ("ADF p-value", f(st.get("ADF_pvalue"), 4)),
            ("KPSS statistic", f(st.get("KPSS_stat"), 4)),
            ("KPSS p-value", f(st.get("KPSS_pvalue"), 4)),
            ("Engine", str(st.get("engine", "—"))),
        ], header=("Test", "Value"))
    else:
        add_kv_table(doc, [
            ("MAE", "[run python backtest.py]"),
            ("RMSE", "[run python backtest.py]"),
            ("MAPE", "[run python backtest.py]"),
            ("Directional Accuracy", "[run python backtest.py]"),
            ("Diebold-Mariano (vs RW)", "[run python backtest.py]"),
            ("Out-of-sample R²", "[run python backtest.py]"),
        ], header=("Metric", "Value"))

    doc.add_heading("9.3 Residual diagnostics", level=2)
    add_bullets(doc, [
        "Residual ACF — chart in dashboard Diagnostics tab.",
        "Histogram with normal fit — chart in dashboard.",
        "Q-Q plot — chart in dashboard.",
        "Ljung-Box test — [FILL IN p-value at lag 10, lag 20].",
    ])

    doc.add_heading("9.4 Failure modes", level=2)
    add_para(
        doc,
        "[Discuss episodes where the model under- or over-predicted significantly. Candidates: "
        "central-bank intervention days, holiday liquidity gaps, sharp basket-currency moves "
        "around major data releases.]",
        italic=True,
    )

    # ------------------------------ 10. Limitations
    doc.add_heading("10. Limitations and Future Work", level=1)
    doc.add_heading("10.1 Known limitations", level=2)
    add_bullets(doc, [
        "Single-regime Kalman — intraday volatility is treated as constant; AM-vs-PM regime shifts are not yet modelled.",
        "Basket coefficients are not individually significant on the full sample under Newey-West HAC inference — BCT's active management of the fixing rate dominates the linear basket signal.",
        "Residual autocorrelation is severe (Ljung-Box p ≈ 0 at lag 10), indicating the AR(1) Kalman alone does not fully whiten residuals.",
        "No exogenous covariates (Brent, DXY, sovereign yield, BCT reserves) — would likely improve fit materially.",
        "Intraday source is yfinance — adequate but not exchange-grade; latency and weekend gaps must be handled.",
        "AM/PM schema exists but historical observations only contain the daily mid; the AM/PM research question can only be answered after several weeks of forward AM/PM collection.",
    ])

    doc.add_heading("10.2 Roadmap", level=2)
    add_bullets(doc, [
        "Vector Error-Correction Model (VECM) benchmark for the IB-Fix cointegration.",
        "Bayesian / particle filter with stochastic volatility on the spread.",
        "Live AM/PM scraping over several weeks to enable the operational AM/PM research question.",
        "Premium / discount signal with hit-ratio backtest as a tradable indicator.",
    ])

    # ------------------------------ 10.3 Markov-switching extension
    doc.add_heading("10.3 Markov-Switching Kalman — preliminary results", level=2)
    add_para(
        doc,
        "As a regime-aware extension to the single-state Kalman, we fit a two-regime "
        "Markov-switching state-space model on the IB-Fix spread. The regimes are "
        "interpreted as a slow-moving liquidity state (quiet vs stressed). Estimation "
        "uses the Hamilton filter combined with Kim's (1994) collapsed Kalman recursion "
        "and joint MLE of all 10 parameters (regime-specific c, φ, Q, R plus the two "
        "diagonal transition probabilities).",
    )
    try:
        import sqlite3
        from clean_returns import load_and_clean
        from msk import fit_msk
        con = sqlite3.connect(str(ROOT / "data" / "tnd.db"))
        _df = load_and_clean(con, lookback_days=10_000)
        con.close()
        if not _df.empty and _df["spread_pub"].dropna().shape[0] >= 50:
            msk = fit_msk(_df["spread_pub"], restarts=2)
            if msk.get("ok"):
                q = msk["regimes"]["quiet"]
                s_ = msk["regimes"]["stressed"]
                add_kv_table(doc, [
                    ("N observations",  str(msk["n_obs"])),
                    ("Log-likelihood",  f"{msk['loglik']:.2f}"),
                    ("AIC",             f"{msk['AIC']:.2f}"),
                    ("BIC",             f"{msk['BIC']:.2f}"),
                    ("Stationary Pr(quiet)", f"{msk['stationary_prob_quiet']:.3f}"),
                    ("Persistence quiet→quiet",       f"{msk['transition']['p_quiet_to_quiet']:.3f}"),
                    ("Persistence stressed→stressed", f"{msk['transition']['p_stressed_to_stressed']:.3f}"),
                ], header=("Field", "Value"))

                add_para(doc, "Regime-conditioned parameters:", bold=True)
                add_kv_table(doc, [
                    ("Quiet — c",       f"{q['c']:+.6f}"),
                    ("Quiet — φ",       f"{q['phi']:+.4f}"),
                    ("Quiet — Q",       f"{q['Q']:.3e}"),
                    ("Quiet — R",       f"{q['R']:.3e}"),
                    ("Stressed — c",    f"{s_['c']:+.6f}"),
                    ("Stressed — φ",    f"{s_['phi']:+.4f}"),
                    ("Stressed — Q",    f"{s_['Q']:.3e}"),
                    ("Stressed — R",    f"{s_['R']:.3e}"),
                ], header=("Parameter", "Value"))
                add_para(
                    doc,
                    "Comparison against the single-regime Kalman MLE log-likelihood "
                    "(see §6.2) gives a likelihood-ratio test of regime presence; the "
                    "AIC and BIC numbers above quantify the trade-off between fit and "
                    "complexity (10 vs 4 parameters).",
                )
            else:
                add_para(doc, f"[MSK fit failed: {msk.get('reason')}]", italic=True)
    except Exception as _e:
        add_para(doc, "[Run python msk.py to populate this section.]", italic=True)

    # ------------------------------ 10.4 Macro overlay
    doc.add_heading("10.4 Macro-overlay regression — exogenous covariates", level=2)
    add_para(
        doc,
        "We test whether global macro covariates carry incremental explanatory "
        "power for the IB-Fix spread, beyond the basket and the AR(1) term. "
        "Three series are pulled from FRED: Brent crude (DCOILBRENTEU), the Fed "
        "trade-weighted broad dollar index (DTWEXBGS), and the CBOE volatility "
        "index VIX (VIXCLS). Four nested specifications are fit; the incremental "
        "R² of each block over the previous one is reported.",
    )
    mac_path = ROOT / "reports" / "macro_overlay.json"
    if mac_path.exists():
        mac = json.loads(mac_path.read_text())
        if mac.get("ok"):
            models = mac["models"]
            add_kv_table(doc, [
                (k, f"k={v['k']}  R²={v['R2']:.4f}  adjR²={v['adj_R2']:.4f}  "
                    f"AIC={v.get('AIC') or 0:.2f}  ΔR²={v['incremental_R2']:+.4f}")
                for k, v in models.items()
            ], header=("Model", "Fit (Newey-West HAC)"))
            add_para(
                doc,
                f"Sample: N={mac['n_obs']}, "
                f"{mac['date_min']} → {mac['date_max']}.",
                italic=True,
            )

            # Detailed M3 coefficient table
            add_para(doc, "Macro coefficients in M3 (HAC SE, two-sided p):", bold=True)
            rows = []
            for c in models["M3_macro"]["coefs"]:
                p = c.get("p")
                sig = (" ***" if p is not None and p < 0.001 else
                       " **"  if p is not None and p < 0.01  else
                       " *"   if p is not None and p < 0.05  else
                       " ."   if p is not None and p < 0.10  else "")
                se_s = f"{c['se']:.5f}" if c.get("se") is not None else "NA"
                t_s  = f"{c['t']:+.3f}"  if c.get("t")  is not None else "NA"
                p_s  = f"{c['p']:.4f}"   if c.get("p")  is not None else "NA"
                rows.append((c["name"],
                             f"est={c['est']:+.6f}  SE={se_s}  t={t_s}  p={p_s}{sig}"))
            add_kv_table(doc, rows, header=("Regressor", "Estimate · SE · t · p"))
        else:
            add_para(doc, f"[Macro overlay not yet runnable: {mac.get('reason')}]", italic=True)
    else:
        add_para(
            doc,
            "[Run python fetch_macro.py followed by python macro_overlay.py "
            "to populate this section.]",
            italic=True,
        )

    # ------------------------------ 11. Conclusion
    doc.add_heading("11. Conclusion", level=1)
    add_para(
        doc,
        "We have delivered a fully working real-time intrinsic-value engine for USD/TND that "
        "satisfies the brief: (i) a basket baseline rigorously estimated by rolling OLS, (ii) a "
        "stochastic liquidity adjustment via an AR(1) Kalman filter on the IB-Fix spread, (iii) "
        "an intraday operating mode that propagates the model state forward between daily IB "
        "observations, and (iv) a complete tooling chain — Excel reporting, Streamlit dashboard, "
        "and reproducible Python code — to operate, monitor, and validate the model end-to-end.",
    )

    # ------------------------------ Appendices
    doc.add_page_break()
    doc.add_heading("Appendix A — Repository Structure", level=1)
    add_code_block(doc, """tnd-fx-pipeline/
├── run_pipeline.py        # daily orchestrator
├── run_realtime.py        # intraday loop runner
├── realtime.py            # tick() — intraday engine
├── fetch_daily.py         # FX APIs + BCT scrape
├── clean_returns.py       # log-returns + spread
├── model.py               # OLS, rolling OLS, Kalman
├── predict.py             # daily prediction
├── export_excel.py        # 3-sheet workbook
├── dashboard.py           # Streamlit UI
├── notify_telegram.py     # optional alerts
├── init_db.py             # SQLite schema + migrations
├── seed_db.py             # CSV bootstrap
├── build_report.py        # this report generator
├── data/
│   ├── tnd.db
│   ├── FX-CLEAN-Data.csv
│   └── ib.csv
└── reports/
    ├── tnd_report_YYYY-MM-DD.xlsx
    └── TND_Intrinsic_Value_Report.docx""")

    doc.add_heading("Appendix B — Database Schema", level=1)
    add_code_block(doc, """CREATE TABLE fx_rates (
    date TEXT PRIMARY KEY, eurusd REAL, gbpusd REAL, usdjpy REAL,
    fix_mid REAL, ib_rate REAL, created_at TEXT
);

CREATE TABLE predictions (
    date TEXT PRIMARY KEY, intrinsic_v1 REAL, intrinsic_v2 REAL,
    w_eurusd REAL, w_gbpusd REAL, w_usdjpy REAL,
    kf_spread REAL, created_at TEXT
);

CREATE TABLE fx_intraday (
    ts TEXT PRIMARY KEY, eurusd REAL, gbpusd REAL, usdjpy REAL,
    source TEXT, created_at TEXT
);

CREATE TABLE intrinsic_intraday (
    ts TEXT PRIMARY KEY, anchor_date TEXT, anchor_fix REAL,
    basket_ret REAL, intrinsic_v1 REAL,
    kf_state REAL, kf_sigma REAL, intrinsic_v2 REAL,
    created_at TEXT
);""")

    doc.add_heading("Appendix C — Key Equations", level=1)
    add_para(doc, "Basket return:")
    add_code_block(doc, "ret(t) = w_const + w_EUR · log(EUR/USD_t / EUR/USD_anchor)\n"
                        "              + w_GBP · log(GBP/USD_t / GBP/USD_anchor)\n"
                        "              + w_JPY · log(USD/JPY_t / USD/JPY_anchor)")
    add_para(doc, "Intrinsic V1 (basket baseline):")
    add_code_block(doc, "intrinsic_v1(t) = fix_mid_anchor · exp( ret(t) )")
    add_para(doc, "Spread Kalman (state-space):")
    add_code_block(doc, "x_t = c + φ · x_{t-1} + w_t,   w_t ~ N(0, Q)\n"
                        "y_t = x_t + v_t,                v_t ~ N(0, R)\n\n"
                        "Predict:  x̂_{t|t-1} = c + φ · x̂_{t-1|t-1}\n"
                        "          P_{t|t-1} = φ² · P_{t-1|t-1} + Q\n"
                        "Update:   K_t = P_{t|t-1} / (P_{t|t-1} + R)\n"
                        "          x̂_{t|t} = x̂_{t|t-1} + K_t · (y_t − x̂_{t|t-1})\n"
                        "          P_{t|t} = (1 − K_t) · P_{t|t-1}")
    add_para(doc, "Intrinsic V2 (full model):")
    add_code_block(doc, "intrinsic_v2(t) = intrinsic_v1(t) + x̂_t")
    add_para(doc, "Premium / Discount (basis points):")
    add_code_block(doc, "premium_bps(t) = (ib_rate(t) − intrinsic_v2(t)) / intrinsic_v2(t) × 10_000")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
